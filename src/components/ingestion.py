"""ingestion.py — lightweight document parsing for DocuMind (Q1 + B2).

Replaces the heavy ``unstructured[all-docs]`` stack with direct, fast parsers:
    - PDF  → PyMuPDF (fitz): per-page text + embedded images
    - DOCX → python-docx: paragraphs + tables + embedded images
    - TXT  → plain read

Text is split on **token** boundaries (Q1) for predictable context size and
cost. Images are extracted here (bytes + page) but **not yet indexed** — the
image-answering step consumes ``parsed["images"]`` to store the real image and
add a vision-model description.
"""

import hashlib
import re
from pathlib import Path
from typing import Callable, List, Tuple

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.logger import get_logger

logger = get_logger(__name__)

# Drop chunks too short to carry meaning, and images too small to be real
# content (icons, bullets, logos) — keeps the index clean and the future
# vision bill down.
_MIN_CHUNK_CHARS = 10
_MIN_IMAGE_BYTES = 3000


# ── Clause/section-aware chunking for legal text ────────────────────────────
#
# A line that STARTS a new clause/section in a regulation: a numbered/lettered
# marker or a heading keyword. Legal text has structure (sections, sub-clauses,
# definitions) that fixed-size token windows cut through mid-clause — hurting
# requirement extraction and retrieval. We split on these markers instead, then
# pack whole clauses up to the token budget.
_CLAUSE_MARKER = re.compile(
    r"^\s*(?:"
    r"\d+\.\d+(?:\.\d+)*[.)]?"                                          # 1.2  1.2.3  1.2.
    r"|\d+[.)]"                                                          # 1.  10)
    r"|\([a-zA-Z0-9]{1,4}\)"                                            # (a) (iv) (12)
    r"|[a-zA-Z][.)]"                                                     # a.  b)  i.
    r"|(?:Section|Chapter|Part|Clause|Paragraph|Rule|Article|Schedule|Annexure|Annex)\b"
    r")\s",
    re.IGNORECASE,
)


def _make_token_counter() -> Callable[[str], int]:
    """A token-length function matching the recursive splitter's (tiktoken gpt2),
    with a rough chars/4 fallback if tiktoken can't load. Used only for the
    packing budget, which is soft — the real token split of an oversized clause
    is delegated to the recursive splitter."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("gpt2")
        return lambda s: len(enc.encode(s))
    except Exception:                                    # pragma: no cover - tiktoken always present
        return lambda s: max(1, len(s) // 4)


def _clause_units(text: str) -> List[str]:
    """Group lines into clause units: each unit starts at a clause/section marker
    line and absorbs the following non-marker lines (its body). Text before the
    first marker (a preamble) is its own leading unit."""
    units: List[str] = []
    cur: List[str] = []
    for line in (text or "").split("\n"):
        if cur and _CLAUSE_MARKER.match(line):
            units.append("\n".join(cur))
            cur = [line]
        else:
            cur.append(line)
    if cur:
        units.append("\n".join(cur))
    return units


def split_legal_text(
    text: str,
    token_splitter: Callable[[str], List[str]],
    count_tokens: Callable[[str], int],
    max_tokens: int,
) -> List[str]:
    """Clause/section-aware split of legal *text*.

    Breaks on clause/section markers, then packs WHOLE clauses up to
    ``max_tokens`` (never cutting a clause across chunks). A single clause larger
    than the budget falls back to ``token_splitter`` (the recursive token
    splitter). No cross-chunk overlap — clause boundaries are the natural split
    points. Text with no markers yields the whole text as one unit, so this
    degrades to plain token-splitting for unstructured input.
    """
    chunks: List[str] = []
    buf: List[str] = []
    buf_tokens = 0
    for unit in _clause_units(text):
        unit_tokens = count_tokens(unit)
        if unit_tokens > max_tokens:                     # oversized clause -> token-split it
            if buf:
                chunks.append("\n".join(buf))
                buf, buf_tokens = [], 0
            chunks.extend(token_splitter(unit))
            continue
        if buf and buf_tokens + unit_tokens > max_tokens:
            chunks.append("\n".join(buf))
            buf, buf_tokens = [], 0
        buf.append(unit)
        buf_tokens += unit_tokens
    if buf:
        chunks.append("\n".join(buf))
    return [c for c in (chunk.strip() for chunk in chunks) if c]


class DocumentProcessor:
    """Parse PDF / DOCX / TXT into token-sized LangChain Documents (+ images)."""

    def __init__(self, config):
        self.config = config
        # from_tiktoken_encoder makes the splitter count *tokens*, not chars.
        self._splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=config.CHUNK_SIZE_TOKENS,
            chunk_overlap=config.CHUNK_OVERLAP_TOKENS,
        )
        self._count_tokens = _make_token_counter()      # for clause-aware packing

    # ── Step 1: parse a file into a normalized dict ───────────────────────

    def process_documents(self, file_path: str) -> dict:
        """Parse *file_path* into a normalized dict, or ``{}`` on failure.

        Returns:
            {
              "filename", "filepath", "filetype",
              "pages":  [(page_number | None, text), ...],
              "images": [{"page_number", "ext", "bytes", "image_hash"}, ...],
            }
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                pages, images, page_images = self._parse_pdf(file_path)
            elif ext == ".docx":
                pages, images, page_images = self._parse_docx(file_path)
            elif ext in (".txt", ".md"):
                pages, images, page_images = self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            logger.debug(
                "Parsed %s: %d page(s), %d image(s) extracted (images deferred to "
                "the image-answering step)", path.name, len(pages), len(images),
            )
            return {
                "filename": path.name,
                "filepath": str(file_path),
                "filetype": ext.lstrip("."),
                "pages": pages,
                "images": images,
                "visual_page_images": page_images,  # {page_no: png_bytes} (B-hybrid)
            }
        except Exception as e:
            logger.error("Error processing %s: %s", file_path, e, exc_info=True)
            return {}

    # ── Step 2: token-chunk the text into LangChain Documents ─────────────

    def build_langchain_documents(self, parsed: dict, clause_aware: bool = False) -> List[Document]:
        """Token-chunk each page's text into LangChain Documents.

        ``clause_aware=True`` uses structure-aware splitting for legal text
        (regulations) — split on clause/section markers, pack whole clauses — so a
        requirement isn't cut mid-clause. Default (False) keeps the plain
        token-window split for policies / general documents.
        """
        if not isinstance(parsed, dict) or not parsed.get("pages"):
            return []

        visual_pages = set(parsed.get("visual_page_images", {}))
        docs: List[Document] = []
        for page_number, text in parsed["pages"]:
            text = (text or "").strip()
            if not text:
                continue
            pieces = (
                split_legal_text(
                    text, self._splitter.split_text, self._count_tokens,
                    self.config.CHUNK_SIZE_TOKENS,
                )
                if clause_aware
                else self._splitter.split_text(text)
            )
            for chunk in pieces:
                chunk = chunk.strip()
                if len(chunk) < _MIN_CHUNK_CHARS:
                    continue
                meta = {
                    "filename": parsed["filename"],
                    "filepath": parsed["filepath"],
                    "filetype": parsed["filetype"],
                    "chunk_type": "text",
                    "source": parsed["filepath"],
                }
                # Only set page_number when we actually have one (PDF). DOCX/TXT
                # have no pages, so leaving it off makes citations read "N/A"
                # instead of a misleading "page 1".
                if page_number is not None:
                    meta["page_number"] = page_number
                    if page_number in visual_pages:
                        meta["has_visual"] = True  # B-hybrid: this page has a snapshot
                docs.append(Document(page_content=chunk, metadata=meta))

        logger.info("Built %d text chunk(s) from %s", len(docs), parsed.get("filename"))
        return docs

    # ── Parsers ───────────────────────────────────────────────────────────

    @staticmethod
    def _image_record(page_number, data: bytes, ext: str) -> dict:
        return {
            "page_number": page_number,
            "ext": ext or "png",
            "bytes": data,
            "image_hash": hashlib.sha256(data).hexdigest(),
        }

    def _parse_pdf(self, file_path) -> Tuple[list, list, dict]:
        pages, images, page_images = [], [], {}
        render = getattr(self.config, "USE_IMAGE_ANSWERING", False)
        dpi = getattr(self.config, "PAGE_IMAGE_DPI", 130)
        with fitz.open(file_path) as doc:
            for page_index, page in enumerate(doc, start=1):
                pages.append((page_index, page.get_text()))
                page_has_image = False
                for img in page.get_images(full=True):
                    xref = img[0]
                    try:
                        base = doc.extract_image(xref)
                    except Exception:
                        continue
                    data = base.get("image", b"")
                    if len(data) >= _MIN_IMAGE_BYTES:
                        images.append(self._image_record(page_index, data, base.get("ext")))
                        page_has_image = True
                # B-hybrid: snapshot pages with visual content so the answer step
                # can show the LLM the real page (text + figures read in place).
                if render and page_has_image:
                    page_images[page_index] = page.get_pixmap(dpi=dpi).tobytes("png")
        return pages, images, page_images

    def _parse_docx(self, file_path) -> Tuple[list, list, dict]:
        doc = DocxDocument(file_path)

        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)

        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    data = rel.target_part.blob
                except Exception:
                    continue
                if len(data) >= _MIN_IMAGE_BYTES:
                    images.append(self._image_record(None, data, "png"))

        # DOCX has no page concept → page_number None (citations show N/A). No
        # page snapshots: DOCX can't be page-rendered with lightweight tools.
        return [(None, text)], images, {}

    def _parse_txt(self, file_path) -> Tuple[list, list, dict]:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        return [(None, text)], [], {}


if __name__ == "__main__":
    from src.components.config import Config

    proc = DocumentProcessor(Config())
    sample = str(
        Path(__file__).parent.parent.parent
        / "docs"
        / "Smart_Signal__Adaptive_Traffic_Signal_Control_using_Reinforcement_Learning_and_Object_Detection.pdf"
    )
    parsed = proc.process_documents(sample)
    docs = proc.build_langchain_documents(parsed)
    print(
        f"pages={len(parsed.get('pages', []))} "
        f"images={len(parsed.get('images', []))} chunks={len(docs)}"
    )
    if docs:
        print("first chunk page:", docs[0].metadata.get("page_number"))
        print(docs[0].page_content[:200])
